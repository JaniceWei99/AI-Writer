"""Convert Markdown text to a .docx file in memory."""

import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(md_text: str, title: str = "") -> io.BytesIO:
    """Parse simplified Markdown and return an in-memory .docx BytesIO buffer."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    for line in md_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue

        # headings
        m = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=level)
            continue

        # unordered list
        m = re.match(r"^[-*+]\s+(.*)", stripped)
        if m:
            doc.add_paragraph(m.group(1), style="List Bullet")
            continue

        # ordered list
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            doc.add_paragraph(m.group(1), style="List Number")
            continue

        # blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            continue

        # horizontal rule
        if re.match(r"^[-*_]{3,}$", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(180, 180, 180)
            continue

        # normal paragraph — apply inline bold / italic
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_inline_runs(paragraph, text: str):
    """Handle **bold** and *italic* inline formatting."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
